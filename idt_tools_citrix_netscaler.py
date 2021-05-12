import datetime
import json
import os

import docx2pdf
import paramiko
import select
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor, Pt
import idt_tools_constant_pm as idtconst
from selenium import webdriver
from selenium.webdriver.common.by import By as WebBy
from selenium.webdriver.support.ui import Select as WebSelect
import urllib.request as urllib2
import time
import idt_tools_word as idtword

cmdX = [
    "vtysh",
    "ter le 0",
    "exit",
    "shell"]


def netscaler_write_command_to_docx(doc, cmd, lines, font_size_cmd, font_size_lines):
    first_line = True
    paragraph = doc.add_paragraph()
    header = paragraph.add_run()
    header.font.size = Pt(font_size_cmd)
    header.font.name = 'Time New Roman'
    header.add_text(cmd)
    header.bold = True
    header.font.color.rgb = RGBColor(0x00, 0x00, 0xff)
    header.font.highlight_color = WD_COLOR_INDEX.YELLOW
    header.add_break()
    a_run = paragraph.add_run()
    header.font.name = 'Consolas'
    a_run.font.size = Pt(font_size_lines)
    for line_idx, line in enumerate(lines):
        if line.strip() != "Done":
            a_run.add_text(line.lstrip())
            a_run.add_break()
        # if line_idx == 0:
        #     p = doc.add_paragraph(line.rstrip())
        # else:
        #     run = p.add_run(line.rstrip())
        #     run.add_break()
        # print(line_idx, line.rstrip())
    # doc.add_paragraph('Hello world!')
    # paraObj1 = doc.add_paragraph('This is a second paragraph.')
    # paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
    # paraObj1.add_run(' This text is being added to the second paragraph.')
    # doc.save('multipleParagraphs.docx')


def cgnat_log_text(device_ip, device_host, device_so, device_type, device_user, device_pw, cmds, xtime):
    # todo 20210422 FrankChen 修改為LOG_YYYYmmDD_HHMM，並調整次目錄順序
    if device_user == "" and device_pw == "":
        print("FRANK no row_id, row_pw used default")
        device_user = idtconst.cgnat_user
        device_pw = idtconst.cgnat_pw
    folder = os.path.join(os.getcwd(), idtconst.str_log, device_type, device_so, device_host)
    if not os.path.exists(folder):
        os.makedirs(folder)
    cgnat_log_text_docx(folder, device_ip, device_host, device_so, device_type, device_user, device_pw, cmds, xtime)
    print("> netscaler_get_https_screenshot")
    netscaler_get_https_screenshot(folder, device_ip, device_host, device_so, device_type, device_user, device_pw)

    # rename Citrix NetScaler - Dashboard.pdf
    pdfs = [f for f in os.listdir(folder) if f[-3:].upper() == "PDF" and f.lower().find("dashboard") > -1]
    if len(pdfs) > 0:
        os.rename(os.path.join(folder, pdfs[0]),
                  os.path.join(folder, "{}_{}_Dashboard.pdf".format(device_ip, device_host)))


def cgnat_log_text_docx(device_log_folder, device_ip, device_host, device_so, device_type, device_user, device_pw, cmds,
                        xtime):
    doc = Document()
    font_size_cmd = 12
    font_size_lines = 9
    if device_user == "" and device_pw == "":
        print("FRANK no row_id, row_pw used default")
        device_user = idtconst.cgnat_user
        device_pw = idtconst.cgnat_pw

    # folder = os.path.join(os.getcwd(), "LOG/%s/%s/" % (device_so, device_type))
    # folder = os.path.join(idtconst.str_log, device_type, device_so)
    folder = device_log_folder
    # docx_file_name = "{}_{}_{}.docx".format(device_ip, xtime, device_host)

    print(cmds)
    ssh = paramiko.SSHClient()
    ssh.load_system_host_keys()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(device_ip,
                username=device_user,
                password=device_pw,
                look_for_keys=False)

    print(cmds)
    third_subject_added = False
    for cmd_idx, cmd in enumerate(cmds):
        print(cmd_idx, cmd)
        ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command(cmd)
        if cmd in cmdX:
            pass
        else:
            output = ssh_stdout.readlines()
            if cmd.lower() == "show host":
                for idx, oline in enumerate(output):
                    print(idx, oline)
                device_host = output[1].strip().split()[-1]
                title = "{}{} 檢查報告".format(device_so, device_host)
                # lines = ["1. 備份系統設定檔"]
                idtword.word_title_to_docx(doc, title, [])
            elif -1 < cmd.lower().find("cat") < cmd.lower().find("ns.conf"):
                a_line = "1. 備份系統設定檔"
                idtword.word_docx_add_highlighted_paragraph_line(doc, a_line, "Time New Roman", 12, True)
            elif -1 < cmd.lower().find("cat") < cmd.lower().find("ZebOS.conf"):
                a_line = "2. 備份路由設定檔"
                idtword.word_docx_add_highlighted_paragraph_line(doc, a_line, "Time New Roman", 12, True)
            elif -1 < cmd.lower().find("messages") < cmd.lower().find("grep"):
                if not third_subject_added:
                    third_subject_added = True
                    a_line = "3. 檢查目前系統 log 訊息"
                    idtword.word_docx_add_highlighted_paragraph_line(doc, a_line, "Time New Roman", 12, True)
            # netscaler_write_command_to_docx(doc, cmd, output, font_size_cmd, font_size_lines)
            idtword.word_write_command_to_docx(doc, cmd, output, font_size_cmd, font_size_lines)
            print(ssh_stderr)

    docx_file_name = "{}({}).docx".format(device_host, device_ip)
    docx_full_path = os.path.join(folder, docx_file_name)
    if not os.path.exists(folder):
        os.makedirs(folder)
    doc.save(docx_full_path)
    ssh.close()
    docx2pdf.convert(docx_full_path)
    if os.path.exists(docx_full_path):
        os.remove(docx_full_path)


def netscaler_show_lsn_session(device_ip, device_host, device_so, device_type, device_user, device_pw, xtime):
    if device_user == "" and device_pw == "":
        print("FRANK no row_id, row_pw used default")
        device_user = idtconst.cgnat_user
        device_pw = idtconst.cgnat_pw

    folder = os.path.join(os.getcwd(), "LOG/%s/%s/" % (device_so, device_type))
    log_file_name = "{}_{}_{}.log".format(device_ip, xtime, device_host)
    log_full_path = os.path.join(folder, log_file_name)
    if not os.path.exists(folder):
        os.makedirs(folder)

    ssh = paramiko.SSHClient()
    ssh.load_system_host_keys()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(device_ip,
                username=device_user,
                password=device_pw,
                look_for_keys=False)
    timeout = 7200
    cmd = "show lsn session"
    print(log_full_path, ssh, cmd)
    myexec(log_full_path, ssh, cmd)
    ssh.close()


def netscaler_show_lsn_session_worker(device_ip, device_host, device_so, device_type, device_user, device_pw, xtime):
    if device_user == "" and device_pw == "":
        print("FRANK no row_id, row_pw used default")
        ID = idtconst.pm_user
        PW = idtconst.pm_pw
    start_time = datetime.datetime.now()
    print(device_so, "start@", start_time)
    folder = os.path.join(os.getcwd(), "LOG/%s/%s/" % (device_so, device_type))
    log_file_name = "{}_{}_{}.log".format(device_ip, xtime, device_host)
    log_full_path = os.path.join(folder, log_file_name)
    if not os.path.exists(folder):
        os.makedirs(folder)
    ssh = paramiko.SSHClient()
    ssh.load_system_host_keys()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(device_ip,
                username=device_user,
                password=device_pw,
                look_for_keys=False)

    cmd = "show lsn session"
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command(cmd)

    worker(log_full_path, ssh_stdout, ssh_stderr)
    ssh.close()
    end_time = datetime.datetime.now()
    print(device_so, "end@", end_time)
    print("total:", end_time - start_time)


def cgnat_show_lsn_session_worker(device_ip, device_host, device_so, wdir, device_user, device_pw, xtime):
    start_time = datetime.datetime.now()
    print(device_so, "start@", start_time)
    folder = wdir
    log_file_name = "{}_{}_{}.log".format(device_ip, xtime, device_host)
    log_full_path = os.path.join(folder, log_file_name)
    if not os.path.exists(folder):
        os.makedirs(folder)
    ssh = paramiko.SSHClient()
    ssh.load_system_host_keys()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(device_ip,
                username=device_user,
                password=device_pw,
                look_for_keys=False)

    cmd = "show lsn session"
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command(cmd)
    worker(log_full_path, ssh_stdout, ssh_stderr, device_host)
    ssh.close()
    end_time = datetime.datetime.now()
    print(device_so, "end@", end_time)
    print("total:", end_time - start_time)


def worker(log_full_path, stdout, stderr, device_host=""):
    # Wait for the command to terminate
    print("Status of worker is {}".format(stdout.channel.exit_status_ready()))
    log_file = open(log_full_path, "w")
    log_file.close()
    count = 0
    while not stdout.channel.exit_status_ready():
        time.sleep(1)
        # print("Status of worker is {}".format(stdout.channel.exit_status_ready()))
        if stdout.channel.recv_ready():
            # Only print data if there is data to read in the channel
            count += 1
            print("{} read and write {} 次".format(device_host, count))
            rl, wl, xl = select.select([stdout.channel], [], [], 0.0)
            if len(rl) > 0:
                # Print data from stdout
                # print("Output: {}".format(stdout.channel.recv(1024).decode("utf8")))
                # buff = stdout.channel.recv(1024).decode("utf8")
                buff = stdout.channel.recv(1048476).decode("utf8")
                log_file = open(log_full_path, "a+")
                log_file.write(buff)
                log_file.close()


def myexec(result_log_file, ssh, cmd, timeout):
    stdin, stdout, stderr = ssh.exec_command(cmd)  # one channel per command
    channel = stdout.channel
    stdin.close()  # we do not need stdin.
    channel.shutdown_write()  # indicate that we're not going to write to that channel anymore

    with open(result_log_file, "w+") as f_log:
        stdout_chunks = []
        # stdout_chunks.append(stdout.channel.recv(len(c.in_buffer)))
        # stdout_chunks.append(stdout.channel.recv(len(stdout.channel.in_buffer)))
        buff = stdout.channel.recv(len(stdout.channel.in_buffer))
        print(buff)
        f_log.write(buff.decode("utf8"))
        # f_log.write(stdout.channel.recv(len(stdout.channel.in_buffer)))
        # chunked read to prevent stalls
        # stop if channel was closed prematurely
        print("before while")
        while not channel.closed:
            print("while")
            readq, _, _ = select.select([stdout.channel], [], [], 0.0)
            print(readq)
            for c in readq:
                if c.recv_ready():
                    # stdout_chunks.append(stdout.channel.recv(len(c.in_buffer)))
                    buff = stdout.channel.recv(len(c.in_buffer))

                    f_log.write(buff)
                    # f_log.write(stdout.channel.recv(len(c.in_buffer)))
                if c.recv_stderr_ready():
                    stderr.channel.recv_stderr(len(c.in_stderr_buffer))
            if stdout.channel.exit_status_ready() and not stderr.channel.recv_stderr_ready() and not stdout.channel.recv_ready():
                stdout.channel.shutdown_read()  # indicate that we're not going to read from this channel anymore
                # close the channel
                stdout.channel.close()
                break  # exit as remote side is finished and our bufferes are empty

        # close all the pseudofiles
        stdout.close()
        stderr.close()

    return True


def myexec_original(ssh, cmd, timeout, want_exitcode=False):
    stdin, stdout, stderr = ssh.exec_command(cmd)  # one channel per command
    # get the shared channel for stdout/stderr/stdin
    channel = stdout.channel

    stdin.close()  # we do not need stdin.
    channel.shutdown_write()  # indicate that we're not going to write to that channel anymore

    # read stdout/stderr in order to prevent read block hangs
    stdout_chunks = []
    # stdout_chunks.append(stdout.channel.recv(len(c.in_buffer)))
    stdout_chunks.append(stdout.channel.recv(len(stdout.channel.in_buffer)))
    # chunked read to prevent stalls
    while not channel.closed:  # stop if channel was closed prematurely
        readq, _, _ = select.select([stdout.channel], [], [], timeout)
        for c in readq:
            if c.recv_ready():
                stdout_chunks.append(stdout.channel.recv(len(c.in_buffer)))
            if c.recv_stderr_ready():
                # make sure to read stderr to prevent stall
                stderr.channel.recv_stderr(len(c.in_stderr_buffer))
        if stdout.channel.exit_status_ready() and not stderr.channel.recv_stderr_ready() and not stdout.channel.recv_ready():
            stdout.channel.shutdown_read()  # indicate that we're not going to read from this channel anymore
            # close the channel
            stdout.channel.close()
            break  # exit as remote side is finished and our bufferes are empty

    # close all the pseudofiles
    stdout.close()
    stderr.close()

    if want_exitcode:
        # exit code is always ready at this point
        return ''.join(stdout_chunks), stdout.channel.recv_exit_status()
    return ''.join(stdout_chunks)


def netscaler_get_https_screenshot(device_log_folder, IP, HOST, SO, DEVICE, device_user, device_pw):
    # RFFOLDER = os.path.join(os.getcwd(), "LOG/%s/%s/%s" % (SO, DEVICE, HOST))
    # RFFOLDER = os.path.join(idtconst.str_log, DEVICE, SO,  HOST)
    RFFOLDER = device_log_folder
    if not os.path.exists(RFFOLDER):
        os.makedirs(RFFOLDER)

    chromedriver = "chromedriver/chromedriver.exe"
    # browser = webdriver.Chrome(chromedriver)
    options = webdriver.ChromeOptions()
    options.add_argument('ignore-certificate-errors')
    # options.add_argument('--allow-outdated-plugins')
    # todo for print to pdf
    chrome_options = webdriver.ChromeOptions()
    appState = {"recentDestinations": [{"id": "Save as PDF", "origin": "local", "account": ""}],
                "selectedDestinationId": "Save as PDF",
                "version": 2,
                "isHeaderFooterEnabled": True,
                "isLandscapeEnabled": False,
                "scalingType": 3,
                "scaling": "72"}
    prefs = {'printing.print_preview_sticky_settings.appState': json.dumps(appState),
             'savefile.default_directory': RFFOLDER,
             'download.default_directory': RFFOLDER}
    options.add_experimental_option('prefs', prefs)
    options.add_argument('kiosk-printing')
    print("prefs", prefs)
    browser = webdriver.Chrome(chromedriver, chrome_options=options)
    URL = "https://{}/menu/st".format(IP)
    print("URL", URL)
    browser.get(URL)
    browser.maximize_window()
    elem = browser.find_element_by_id("username")
    elem.send_keys(device_user)
    elem = browser.find_element_by_name("password")
    elem.send_keys(device_pw)
    btn = browser.find_element_by_class_name('login_button')
    # print(btn.text)
    # btn = browser.find_element_by_id("url")
    # print(btn.text)
    btn.click()
    time.sleep(20)
    dashboard1 = "%s/dashboard1.png" % (RFFOLDER)
    browser.get_screenshot_as_file(dashboard1)
    # dashboard2 = "%s/dashboard2.png" % (RFFOLDER)
    # browser.find_element_by_class_name('ns_body').screenshot(dashboard2)
    browser.execute_script('window.print();')
    browser.close()
    browser.quit()


def netscaler_show_lsn_client(device_ip, device_user, device_pw):
    if device_user == "" and device_pw == "":
        print("FRANK no row_id, row_pw used default")
        ID = idtconst.pm_user
        PW = idtconst.pm_pw

    folder = os.path.join(os.getcwd(), "LOG/%s/" % (device_ip))
    log_file_name = "LSN_{}.csv"
    log_full_path = os.path.join(folder, log_file_name)
    if not os.path.exists(folder):
        os.makedirs(folder)

    ssh = paramiko.SSHClient()
    ssh.load_system_host_keys()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(device_ip,
                username=device_user,
                password=device_pw,
                look_for_keys=False)
    timeout = 7200
    cmd1 = "sh lsn client"
    cmd2 = "show lsn deterministicNat -clientname {}"
    ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command(cmd1)
    clients = ssh_stdout.readlines()
    print(clients)
    for client in clients:
        client_info = client.split()
        if len(client_info) > 1:
            cmd_nat = cmd2.format(client_info[-1])
            ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command(cmd_nat)
            client_nats = ssh_stdout.readlines()
            client_csv = log_file_name.format(client_info[-1])
            client_csv_full = os.path.join(folder, client_csv)
            with open(client_csv_full, mode="a") as client_csv_file:
                client_csv_file.writelines(client_nats)

    ssh.close()

# appState = {"recentDestinations": [{"id": "Save as PDF", "origin": "local", "account": ""}],
#                 "selectedDestinationId": "Save as PDF",
#                 "version": 2,
#                 "isHeaderFooterEnabled": False,
#                 "isLandscapeEnabled": False,
#                 "scalingType": 3,
#                 "scaling": "141"}


# def OLD_netscaler_get_https_screenshot(IP, HOST, SO, DEVICE, device_user, device_pw):
#     RFFOLDER = "LOG/%s/%s/%s" % (SO, DEVICE, HOST)
#     if not os.path.exists(RFFOLDER):
#         os.makedirs(RFFOLDER)
#
#     chromedriver = "chromedriver/chromedriver.exe"
#     # browser = webdriver.Chrome(chromedriver)
#     options = webdriver.ChromeOptions()
#     options.add_argument('ignore-certificate-errors')
#     # options.add_argument('--allow-outdated-plugins')
#     # for print to pdf
#     settings = {
#         "recentDestinations": [{
#             "id": "Save as PDF",
#             "origin": "local",
#             "account": "",
#         }],
#         "selectedDestinationId": "Save as PDF",
#         "version": 2
#     }
#     prefs = {'printing.print_preview_sticky_settings.appState': json.dumps(settings)}
#
#     options.add_experimental_option('prefs', prefs)
#     options.add_argument('--kiosk-printing')
#
#     browser = webdriver.Chrome(chromedriver, chrome_options=options)
#     URL = "https://{}/menu/st".format(IP)
#     browser.get(URL)
#     browser.maximize_window()
#     elem = browser.find_element_by_id("username")
#     elem.send_keys(device_user)
#     elem = browser.find_element_by_name("password")
#     elem.send_keys(device_pw)
#     btn = browser.find_element_by_class_name('login_button')
#     # print(btn.text)
#     # btn = browser.find_element_by_id("url")
#     # print(btn.text)
#     btn.click()
#     time.sleep(20)
#     dashboard1 = "%s/dashboard1.png" % (RFFOLDER)
#     browser.get_screenshot_as_file(dashboard1)
#     # dashboard2 = "%s/dashboard2.png" % (RFFOLDER)
#     # browser.find_element_by_class_name('ns_body').screenshot(dashboard2)
#     browser.execute_script('window.print();')
#
#     browser.close()
#     browser.quit()
