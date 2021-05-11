import os
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor, Pt
import docx2pdf

# todo before 20210510 Frank
#  1. to write word document (docx)
#     pip install python-docx
#     import docx
#  x
#  20210510 Frank
#  如果 pyinstaller 封裝後無法執行(閃退) docx2pdf的問題請
#  1. 開啟或產生hook-docx2pdf.py， ~\venv\Lib\site-packages\PyInstaller\hooks\hook-docx2pdf.py
#  2. 將檔案內容改為如下:
#     from PyInstaller.utils.hooks import collect_all
#     datas, binaries, hiddenimports = collect_all('docx2pdf')
#  3. 如仍有問題，至BUILD目錄，清(移)除與此專案相關的目錄


def word_title_to_docx(doc, title, lines):
    first_line = True
    paragraph = doc.add_paragraph()
    header = paragraph.add_run()
    header.add_text(title)
    header.bold = True
    header.font.color.rgb = RGBColor(0x00, 0x00, 0xff)
    header.font.size = Pt(20)
    header.font.highlight_color = WD_COLOR_INDEX.YELLOW
    header.add_break()
    a_run = paragraph.add_run()
    a_run.font.size = Pt(16)
    a_run.font.name = 'Courier New'
    for line_idx, line in enumerate(lines):
        if line.strip() != "Done":
            a_run.add_text(line.lstrip())
            a_run.add_break()


def word_docx_add_highlighted_paragraph_line(doc, a_line, font_name, font_size, is_bold):
    first_line = True
    paragraph = doc.add_paragraph()
    header = paragraph.add_run()
    header.add_text(a_line)
    header.bold = is_bold
    header.font.color.rgb = RGBColor(0x00, 0x00, 0xff)
    header.font.size = Pt(font_size)
    header.font.name = font_name
    header.font.highlight_color = WD_COLOR_INDEX.YELLOW


def word_write_command_to_docx(doc, cmd, lines):
    first_line = True
    paragraph = doc.add_paragraph()
    header = paragraph.add_run()
    header.add_text(cmd)
    header.bold = True
    header.font.size = Pt(10)
    header.font.name = 'Times New Roman'
    header.font.color.rgb = RGBColor(0x00, 0x00, 0xff)
    header.font.highlight_color = WD_COLOR_INDEX.YELLOW
    header.add_break()
    a_run = paragraph.add_run()
    a_run.font.size = Pt(8)
    # a_run.font.name = 'Times New Roman'
    a_run.font.name = 'Consolas'
    for line_idx, line in enumerate(lines):
        if line.strip() != "Done":
            a_run.add_text(line.lstrip())
            # if line_idx < len(lines)-1:
            a_run.add_break()


def word_log_txt_file_to_docx(log_file_full, title):
    doc = Document()
    word_docx_add_highlighted_paragraph_line(doc, title, 'Times New Roman', 10, True)
    docx_file_full = log_file_full[:-4] + ".docx"
    with open(log_file_full, mode="r") as log_file_obj:
        txt_lines = log_file_obj.readlines()
    start_cmd = False
    cmd = ""
    cmd_lines = []
    for txt_line in txt_lines:
        # print(txt_line)
        if len(txt_line.strip()) == 0:
            pass
        elif txt_line.upper().find("#SH") > 0 or txt_line.upper().find("#ADMIN") > 0:
            if len(cmd_lines) > 0:
                # write out the previous cmd
                word_write_command_to_docx(doc, cmd, cmd_lines)
            start_cmd = True
            cmd = txt_line
            cmd_lines = []
        elif start_cmd:
            cmd_lines.append(txt_line)
            if len(cmd_lines) > 15:
                # write out this cmd
                word_write_command_to_docx(doc, cmd, cmd_lines)
                start_cmd = False
                cmd_lines = []
                cmd = ""
    doc.save(docx_file_full)
    docx2pdf.convert(docx_file_full)
    if os.path.exists(docx_file_full):
        os.remove(docx_file_full)


def word_log_txt_file_to_docx_cmts(log_file_full, title):
    doc = Document()
    word_docx_add_highlighted_paragraph_line(doc, title, 'Times New Roman', 10, True)
    docx_file_full = log_file_full[:-4] + ".docx"
    with open(log_file_full, mode="r") as log_file_obj:
        txt_lines = log_file_obj.readlines()
    start_cmd = False
    cmd = ""
    cmd_lines = []
    for txt_line in txt_lines:
        print(txt_line)
        if len(txt_line.strip()) == 0:
            pass
        elif start_cmd:
            if txt_line.upper().find("#SH") > 0:
                start_cmd = False
                word_write_command_to_docx(doc, cmd, cmd_lines)
                break
            else:
                cmd_lines.append(txt_line)

        elif txt_line.upper().find("#SH") > 0 and txt_line.upper().find("VERSION") > 0:
            start_cmd = True
            cmd = txt_line
            cmd_lines = []

    doc.save(docx_file_full)
    docx2pdf.convert(docx_file_full)
    if os.path.exists(docx_file_full):
        os.remove(docx_file_full)


if __name__ == "main":
    txt1 = r"C:\PyDEV\KBRO\PM\LOG_2021_Q2\ASRN9K\16_FM\ASR\FOX1729GWHJ_FM-I-ASR-02.txt"
    txt2 = r"C:\PyDEV\KBRO\PM\LOG_2021_Q2\ASRN9K\16_FM\ASR\FOX1730GDJ2_FM-I-ASR-01.txt"
    # title = "{} {}檢查報告"
    title1 = "{} {}檢查報告".format("豐盟", "FOX1729GWHJ_FM-I-ASR-02")
    title2 = "{} {}檢查報告".format("豐盟", "FOX1730GDJ2_FM-I-ASR-01")
    word_log_txt_file_to_docx(txt1, title1)
    word_log_txt_file_to_docx(txt2, title2)

