import datetime
import json
import os
import traceback

import paramiko
import select
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
from selenium.webdriver import ActionChains
from selenium.webdriver.common.keys import Keys

import idt_tools_constant_pm as idtconst
from selenium import webdriver
from selenium.webdriver.common.by import By as WebBy
from selenium.webdriver.support.ui import Select as WebSelect
import urllib.request as urllib2
import time
import sys

cmdX = [
        "vtysh",
        "ter le 0",
        "exit",
        "shell"]

def dell_qb_log_and_screens(device_ip, device_host, device_so, device_type, device_user, device_pw, cmds, xtime):
    device_ip1 = device_ip.split(",")[0].strip()
    device_ip2 = device_ip.split(",")[-1].strip()
    folder = os.path.join(os.getcwd(), idtconst.str_log, device_type, device_so,  device_host)
    if not os.path.exists(folder):
        os.makedirs(folder)
    # output to document
    netscaler_log_text_docx(folder, device_ip1, device_host, device_so, device_type, device_user, device_pw, cmds, xtime)
    # output to log
    qb_dell_log_and_config_txts(folder, device_ip1, device_host, device_so, device_type, device_user,
                                device_pw, cmds, xtime)

    qb_get_https_screenshot(folder, device_ip1, device_host, device_so, device_type, device_user, device_pw)
    device_idrac_user = "idtech"
    device_idrac_pw = "Idtech123!"
    qb_get_idrac_screenshot(folder, device_ip2, device_host, device_so, device_type, device_idrac_user, device_idrac_pw)
    # # rename Citrix NetScaler - Dashboard.pdf
    # pdfs = [f for f in os.listdir(folder) if f[-3:].upper() == "PDF" and f.lower().find("dashboard") > -1]
    # if len(pdfs) > 0:
    #     os.rename(os.path.join(folder, pdfs[0]), os.path.join(folder, "{}_{}_Dashboard.pdf".format(device_ip, device_host)))


def netscaler_write_command_to_docx(doc, cmd, lines):
    first_line = True
    paragraph = doc.add_paragraph()
    header = paragraph.add_run()
    header.add_text(cmd)
    header.bold = True
    header.font.color.rgb = RGBColor(0x00, 0x00, 0xff)
    header.font.highlight_color = WD_COLOR_INDEX.YELLOW
    header.add_break()
    a_run = paragraph.add_run()
    for line_idx, line in enumerate(lines):
        if line.strip() != "Done":
            a_run.add_text(line.lstrip())
            a_run.add_break()
        # if line_idx == 0:
        #     p = doc.add_paragraph(line.rstrip())
        # else:
        #     run = p.add_run(line.rstrip())
        #     run.add_break()
        # print(line_idx, line.rstrip())
    # doc.add_paragraph('Hello world!')
    # paraObj1 = doc.add_paragraph('This is a second paragraph.')
    # paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
    # paraObj1.add_run(' This text is being added to the second paragraph.')
    # doc.save('multipleParagraphs.docx')


def write_command_to_txt(file_name, cmd, lines):
    write_start = False
    if not os.path.exists(file_name):
        write_start = True
    with open(file_name, mode="a") as file_obj:
        if write_start:
            file_obj.write("[START] @ " + datetime.datetime.now().strftime("%Y/%m/%d %H:%M:%S") + "\n")
        file_obj.write("\n" + "#" + cmd + "\n")
        for line_idx, line in enumerate(lines):
            if len(line.strip()) > 0:
                file_obj.write(line.rstrip() + "\n")


def qb_dell_log_and_config_txts(device_log_folder, device_ip, device_host, device_so, device_type, device_user, device_pw, cmds, xtime):
    print("qb_dell_log_and_config_txts(", device_log_folder, device_ip, device_host, device_so, device_type, device_user,
                                device_pw, cmds, xtime, ")")
    if device_user == "" and device_pw == "":
        print("FRANK no row_id, row_pw used default")
        if device_type == "QB":
            device_user = idtconst.qb_user
            device_pw = idtconst.qb_pw
        else:
            device_user = idtconst.pm_user
            device_pw = idtconst.pm_pw

    folder = device_log_folder

    log_file_name = "{}_{}.txt".format(device_ip, device_host)
    config_file_name = "{}_{}_config.txt".format(device_ip, device_host)
    log_full_path = os.path.join(folder, log_file_name)
    config_full_path = os.path.join(folder, config_file_name)
    if not os.path.exists(folder):
        os.makedirs(folder)
    print(cmds)
    ssh = paramiko.SSHClient()
    ssh.load_system_host_keys()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(device_ip,
                username=device_user,
                password=device_pw,
                look_for_keys=False)

    print(cmds)
    for cmd_idx, cmd in enumerate(cmds):
        print(cmd_idx, cmd)
        ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command(cmd)
        if cmd in cmdX:
            pass
        else:
            output = ssh_stdout.readlines()
            if cmd == "show running-config | nomore":
                write_command_to_txt(config_full_path, cmd, output)
            else:
                write_command_to_txt(log_full_path, cmd, output)
            print(ssh_stderr.readlines())
    ssh.close()


def netscaler_log_text_docx(device_log_folder, device_ip, device_host, device_so, device_type, device_user, device_pw, cmds, xtime):
    doc = Document()
    if device_user == "" and device_pw == "":
        print("FRANK no row_id, row_pw used default")
        if device_type == "QB":
            ID = idtconst.qb_user
            PW = idtconst.qb_pw
        else:
            ID = idtconst.pm_user
            PW = idtconst.pm_pw

    folder = device_log_folder
    # docx_file_name = "{}_{}_{}.docx".format(device_ip, xtime, device_host)
    docx_file_name = "{}_{}.docx".format(device_ip, device_host)
    docx_full_path = os.path.join(folder, docx_file_name)
    if not os.path.exists(folder):
        os.makedirs(folder)
    print(cmds)
    ssh = paramiko.SSHClient()
    ssh.load_system_host_keys()
    ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
    ssh.connect(device_ip,
                username=device_user,
                password=device_pw,
                look_for_keys=False)

    print(cmds)
    for cmd_idx, cmd in enumerate(cmds):
        print(cmd_idx, cmd)
        ssh_stdin, ssh_stdout, ssh_stderr = ssh.exec_command(cmd)
        if cmd in cmdX:
            pass
        else:
            output = ssh_stdout.readlines()
            # doc.add_heading(text=cmd, level=1)
            # header = doc.add_paragraph(cmd)
            netscaler_write_command_to_docx(doc, cmd, output)
            # print(ssh_stderr.readlines())
            doc.save(docx_full_path)
    ssh.close()


def qb_get_https_screenshot(device_log_folder, IP, HOST, SO, DEVICE, device_user, device_pw):
    # RFFOLDER = os.path.join(os.getcwd(), "LOG/%s/%s/%s" % (SO, DEVICE, HOST))
    # RFFOLDER = os.path.join(idtconst.str_log, DEVICE, SO,  HOST)
    RFFOLDER = device_log_folder
    if not os.path.exists(RFFOLDER):
        os.makedirs(RFFOLDER)

    chromedriver = "chromedriver/chromedriver.exe"
    # browser = webdriver.Chrome(chromedriver)
    options = webdriver.ChromeOptions()
    options.add_argument('ignore-certificate-errors')
    # options.add_argument('--allow-outdated-plugins')
    # todo for print to pdf
    chrome_options = webdriver.ChromeOptions()
    appState = {"recentDestinations": [{"id": "Save as PDF", "origin": "local", "account": ""}],
                "selectedDestinationId": "Save as PDF",
                "version": 2,
                "isHeaderFooterEnabled": True,
                "isLandscapeEnabled": False,
                "scalingType": 3,
                "scaling": "72"}
    prefs = {'printing.print_preview_sticky_settings.appState': json.dumps(appState),
             'savefile.default_directory': RFFOLDER,
             'download.default_directory': RFFOLDER}
    options.add_experimental_option('prefs', prefs)
    options.add_argument('kiosk-printing')

    browser = webdriver.Chrome(chromedriver, chrome_options=options)
    # URL = "https://{}/menu/st".format(IP)
    URL = "http://{}/login/client/login/?next=/m/client/system/&appId=mng".format(IP)
    browser.get(URL)
    browser.maximize_window()
    elem = browser.find_element_by_id("username")
    elem.send_keys(device_user)
    elem = browser.find_element_by_name("password")
    elem.send_keys(device_pw)
    btn = browser.find_element_by_class_name('loginButton')
    # print(btn.text)
    # btn = browser.find_element_by_id("url")
    # print(btn.text)
    btn.click()
    time.sleep(10)
    dashboard1 = "{}/{}_{}_viewer.png".format(RFFOLDER, IP, HOST)
    browser.get_screenshot_as_file(dashboard1)
    # dashboard2 = "%s/dashboard2.png" % (RFFOLDER)
    # browser.find_element_by_class_name('ns_body').screenshot(dashboard2)
    # browser.execute_script('window.print();')
    browser.close()
    browser.quit()


def qb_get_idrac_screenshot(device_log_folder, IP, HOST, SO, DEVICE, device_user, device_pw):
    try:
        RFFOLDER = device_log_folder
        if not os.path.exists(RFFOLDER):
            os.makedirs(RFFOLDER)

        chromedriver = "chromedriver/chromedriver.exe"
        # browser = webdriver.Chrome(chromedriver)
        options = webdriver.ChromeOptions()
        options.add_argument('ignore-certificate-errors')
        # options.add_argument('--allow-outdated-plugins')
        # todo for print to pdf
        chrome_options = webdriver.ChromeOptions()
        appState = {"recentDestinations": [{"id": "Save as PDF", "origin": "local", "account": ""}],
                    "selectedDestinationId": "Save as PDF",
                    "version": 2,
                    "isHeaderFooterEnabled": True,
                    "isLandscapeEnabled": False,
                    "scalingType": 3, "scaling": "72"}
        prefs = {'printing.print_preview_sticky_settings.appState': json.dumps(appState),
                 'savefile.default_directory': RFFOLDER,
                 'download.default_directory': RFFOLDER}
        options.add_experimental_option('prefs', prefs)
        options.add_argument('kiosk-printing')
        # options.add_argument("--lang=en")

        browser = webdriver.Chrome(chromedriver, chrome_options=options)
        # URL = "https://{}/menu/st".format(IP)
        URL = "https://{}/login.html".format(IP)
        browser.get(URL)
        browser.maximize_window()
        time.sleep(5)
        elem_user = browser.find_element_by_id("user")

        actionChains = ActionChains(browser)
        actionChains.move_to_element(elem_user).click().perform()
        # actionChains.move_to_element(elem_user).send_keys(device_user, Keys.RETURN).perform()
        ActionChains(browser).move_to_element(elem_user).send_keys(device_user).perform()
        print("user key sent")
        elem_pw = browser.find_element_by_name("password")
        print("found elem_pw")
        actionChains_pw = ActionChains(browser)
        actionChains_pw.move_to_element(elem_pw).click().perform()
        actionChains_pw.move_to_element(elem_pw).send_keys(device_pw, Keys.RETURN).perform()
        print("pw key sent")

        time.sleep(25)
        dashboard1 = "{}/{}_{}_idrac.png".format(RFFOLDER, IP, HOST)
        browser.get_screenshot_as_file(dashboard1)
        # dashboard2 = "%s/dashboard2.png" % (RFFOLDER)
        # browser.find_element_by_class_name('ns_body').screenshot(dashboard2)
        # browser.execute_script('window.print();')
        browser.close()
        browser.quit()
    except Exception as ex:
        error_class = ex.__class__.__name__  # 取得錯誤類型
        detail = ex.args[0]  # 取得詳細內容
        cl, exc, tb = sys.exc_info()  # 取得Call Stack
        lastCallStack = traceback.extract_tb(tb)[-1]  # 取得Call Stack的最後一筆資料
        print("detail:", detail)
        print("cl:", cl)
        print("exc:", exc)
        print("tb:", tb)
        print("lastCallStack:", lastCallStack)